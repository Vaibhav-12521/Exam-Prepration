"""
Convert the 4 Image Processing study documents from markdown to .docx.
Uses the same DCCN-style naming convention (UPPERCASE filenames).
"""
import re
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(r"C:\Users\vaibh\Desktop\SRMU\exam Prepration\Image Processing\exam_output")

DOCS = {
    "SYLLABUS_NOTES.md": ("SYLLABUS_NOTES.docx", "IMAGE PROCESSING — SYLLABUS NOTES",
                          "Course: BCA / B.Sc(IT) | Subject Code: UCS6105 | Sem VI"),
    "SOLUTIONS.md": ("SOLUTIONS.docx", "IMAGE PROCESSING — PROBLEM SET SOLUTIONS",
                     "Course: BCA / B.Sc(IT) | Subject Code: UCS6105 | Sem VI"),
    "IMPORTANT_QUESTIONS.md": ("IMPORTANT_QUESTIONS.docx", "IMAGE PROCESSING — IMPORTANT QUESTIONS",
                               "Course: BCA / B.Sc(IT) | Subject Code: UCS6105 | Sem VI"),
    "REVISION_SHEET.md": ("REVISION_SHEET.docx", "IMAGE PROCESSING — LAST-MINUTE REVISION",
                          "Course: BCA / B.Sc(IT) | Subject Code: UCS6105 | Sem VI"),
    "MOST_IMPORTANT_TOPICS.md": ("MOST_IMPORTANT_TOPICS.docx", "IMAGE PROCESSING — 28 MOST IMPORTANT TOPICS",
                                 "Course: BCA / B.Sc(IT) | Subject Code: UCS6105 | Sem VI"),
}


INLINE_RE = re.compile(
    r"(\*\*([^*]+?)\*\*)"   # bold
    r"|(`([^`]+?)`)"        # inline code
    r"|(\*([^*\n]+?)\*)"    # italic
)


def add_runs(paragraph, text):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1):
            r = paragraph.add_run(m.group(2))
            r.bold = True
        elif m.group(3):
            r = paragraph.add_run(m.group(4))
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif m.group(5):
            r = paragraph.add_run(m.group(6))
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_page_number(footer_paragraph):
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def make_title_page(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n\n\n" + title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("\n" + subtitle)
    r2.font.size = Pt(14)
    r2.italic = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("\n\nFinal Semester VI Examination")
    r3.font.size = Pt(13)
    r3.bold = True

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("\nShri Ramswaroop Memorial University")
    r4.font.size = Pt(12)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run("\n\nDate: " + date.today().strftime("%d %B %Y"))
    r5.font.size = Pt(11)

    doc.add_page_break()


def add_table_from_md(doc, header_cells, rows):
    n_cols = len(header_cells)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_cells):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        add_runs(p, h.strip())
        for run in p.runs:
            run.bold = True

    for r_i, row in enumerate(rows, start=1):
        for c_i, cell_text in enumerate(row[:n_cols]):
            tc = table.rows[r_i].cells[c_i]
            tc.text = ""
            p = tc.paragraphs[0]
            add_runs(p, cell_text.strip())

    doc.add_paragraph()


def parse_table_block(lines, start):
    if "|" not in lines[start]:
        return None, start
    if start + 1 >= len(lines):
        return None, start
    sep = lines[start + 1]
    if not re.match(r"^\s*\|?\s*[:\-]+\s*(\|\s*[:\-]+\s*)+\|?\s*$", sep):
        return None, start

    def split_row(row):
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        return cells

    header = split_row(lines[start])
    rows = []
    i = start + 2
    while i < len(lines):
        row = lines[i]
        if "|" not in row or row.strip() == "":
            break
        rows.append(split_row(row))
        i += 1
    return (header, rows), i


def convert_md(md_path: Path, docx_path: Path, title: str, subtitle: str):
    text = md_path.read_text(encoding="utf-8", errors="replace")
    lines = text.split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = doc.sections[0]
    footer = section.footer
    footer_p = footer.paragraphs[0]
    add_page_number(footer_p)

    make_title_page(doc, title, subtitle)

    in_code = False
    code_buf = []
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buf))
                r.font.name = "Consolas"
                r.font.size = Pt(9)
                in_code = False
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        if stripped in ("---", "***", "___"):
            doc.add_paragraph()
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        table_data, new_i = parse_table_block(lines, i)
        if table_data is not None:
            header, rows = table_data
            add_table_from_md(doc, header, rows)
            i = new_i
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            level = min(level, 4)
            h = doc.add_heading(level=level)
            add_runs(h, heading_text)
            i += 1
            continue

        if re.match(r"^\s*[-*+]\s+", line):
            text_part = re.sub(r"^\s*[-*+]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, text_part)
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            text_part = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_runs(p, text_part)
            i += 1
            continue

        if stripped.startswith(">"):
            text_part = stripped.lstrip("> ").rstrip()
            p = doc.add_paragraph()
            r = p.add_run(text_part)
            r.italic = True
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    doc.save(docx_path)
    print(f"  Saved {docx_path.name} ({docx_path.stat().st_size // 1024} KB)")


def main():
    print(f"Converting markdown to docx in: {HERE}")
    for md_name, (docx_name, title, subtitle) in DOCS.items():
        md_path = HERE / md_name
        docx_path = HERE / docx_name
        if not md_path.exists():
            print(f"  Missing: {md_name}")
            continue
        print(f"-> {md_name}")
        convert_md(md_path, docx_path, title, subtitle)
    print("Done.")


if __name__ == "__main__":
    main()
