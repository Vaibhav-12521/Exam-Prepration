"""Extract text from docx for reading."""
from docx import Document
from pathlib import Path

path = Path(r"C:\Users\vaibh\Desktop\SRMU\exam Prepration\Machine learning\IMPOT\Departmental Internal Test2(ML) all set.docx")
doc = Document(path)

for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

for table in doc.tables:
    print("\n--- TABLE ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
    print("--- END TABLE ---\n")
